#!/usr/bin/python
#coding:utf-8
# tested in Win
# Version: 20190507


import time,os,sys
from docx import Document
from docx.shared import Inches
import comtypes.client

"""
Serial Number:          6CU830G700
Site:   SHANGHAI
Area/Building/Room:
IDX2/A24/GSHFRB024A2/Bay9
"""


def sir():
    wp = os.path.dirname(os.path.realpath(__file__))
    # os.chdir(wp)
    name = input('Server Name >> ')
    SN = input('Serial Number >> ')
    site = input('Press Enter for default Shanghai >> ')
    if site == '': site = 'Shanghai'
    # loc = input('eg:IDX2/A24/GSHFRB024A2/Bay7 >>> ')
    DC = input('DataCenter default IDX2 >>')
    if DC == '': DC = 'IDX2'
    Rack = input('Rack >>')
    if Rack == '' : raise 'Rack info missing'
    loc = f'{DC}/{Rack}'

    Enclosure = input('Enclosure >>')
    Bay = input('Bay number >>')
    if Enclosure != '' and Bay != '':
        loc = f'{loc}/{Enclosure}/Bay {Bay}'
    

    model_dict = {
        'bl9':('813198-B21','ProLiant BL460c Gen9'),
        'bl10':('','ProLiant BL460c Gen10'),
        'dl10':('10','10')
    }

    model = input('bl9/dl10 ?>>')
    if model == 'bl9':
        m = model_dict['bl9']
    elif model == 'dl10':
        m = model_dict['dl10']
    else:
        pass

    SIR = os.path.join(wp,'SIR.docx')
    out_file = os.path.join(wp,'%s.pdf' % name)


    doc = Document(SIR)
    tables = doc.tables

    tinfo = tables[0]
    cName = tinfo.cell(0,0) # (row,column)
    cName.text = 'System Name:      '+name.upper()
    cSN = tinfo.cell(0,1)
    cSN.text = 'Serial Number:          '+SN.upper()
    cSite = tinfo.cell(1,0)
    cSite.text = 'Site:   '+site.upper()
    cloc = tinfo.cell(1,1)
    cloc.text = 'Area/Building/Room:\n'+loc.upper()

    tmodel = tables[1]
    cProduct = tmodel.cell(1,0)
    cProduct.text = m[0]
    cDescription = tmodel.cell(1,2)
    cDescription.text = m[1]

    tsign = tables[3]
    cdate = tsign.cell(1,1)
    cdate.text = time.strftime('%Y-%m-%d',time.localtime(time.time()))+'\n'

    # sig = r'E:\UT\signature.png'
    # s = t1.cell(1,2).add_paragraph()
    # r = s.add_run()
    # r.add_picture(sig,width = Inches(0.5), height = Inches(0.5) )

    try:
        doc.save(SIR)
    except PermissionError as e:
        print(e)
        print('Is file being opened?')

    # os.startfile(SIR,'print')

    wdFormatPDF = 17
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(SIR)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

if __name__=='__main__':
    try:
        sir()
    except KeyboardInterrupt:
        print('ctrl + c')

