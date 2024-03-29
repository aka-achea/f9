#!/usr/bin/python3
#coding:utf-8
# tested in Win

__Version__ = 20211230


import time,os,sys,argparse
from docx import Document
from docx.shared import Inches
import comtypes.client
import openpyxl 

"""
Serial Number:          6CU830G700
Site:   SHANGHAI
Area/Building/Room:
IDX2/A24/GSHFRB024A2/Bay9
"""

# wp = os.path.dirname(os.path.realpath(__file__))
wp = r'E:\ptool\SIR'

model_dict = {
    '460g9':('813198-B21','ProLiant BL460 Gen9'),
    '460g10':('863442-B21','ProLiant BL460 Gen10'),
    '380g10':('868703-B21','ProLiant DL380 Gen10'),
    '580g10':('869854-B21','ProLiant DL580 Gen10'),
    'nutanix':('P18230-B21',"ProLiant DX360 Gen10"),
    's480g10':('871942-B21','HPE Synergy 480 Gen10')
}

def sir(srv_dict:dict):
    '''fill SIR form'''
    SIR = os.path.join(wp,'SIR.docx')
    out_file = os.path.join(wp,'%s.pdf' % srv_dict['name'])

    doc = Document(SIR)
    tables = doc.tables

    tinfo = tables[0]
    cName = tinfo.cell(0,0) # (row,column)
    cName.text = 'System Name:      '+srv_dict['name'].upper()
    cSN = tinfo.cell(0,1)
    cSN.text = 'Serial Number:          '+srv_dict['SN'].upper()
    cSite = tinfo.cell(1,0)
    cSite.text = 'Site:   '+srv_dict['site'].upper()
    cloc = tinfo.cell(1,1)
    cloc.text = 'Area/Building/Room:\n'+srv_dict['loc'].upper()

    tmodel = tables[1]
    cProduct = tmodel.cell(1,0)
    cProduct.text = srv_dict['model'][0]
    cDescription = tmodel.cell(1,2)
    cDescription.text = srv_dict['model'][1]

    tsign = tables[3]
    cdate = tsign.cell(1,1)
    cdate.text = '\n'+time.strftime('%Y-%m-%d',time.localtime(time.time()))+'\n'

    # sig = r'E:\UT\signature.png'
    # s = t1.cell(1,2).add_paragraph()
    # r = s.add_run()
    # r.add_picture(sig,width = Inches(0.5), height = Inches(0.5) )

    try:
        doc.save(SIR)
    except PermissionError as e:
        print(e)
        print('Is file being opened?')

    # os.startfile(SIR,'print')
    try:
        wdFormatPDF = 17
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(SIR)
        doc.SaveAs(out_file, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()
    except :
        # print(e)
        print(f"{srv_dict['name']} already created")


def select():
    '''interact input'''
    opt = """
    1. ProLiant BL460c Gen10
    2. ProLiant DL380 Gen10
    3. ProLiant BL460c Gen9
    """
    print(opt)

    m = input('default choice 1 >>>>')
    if m == '3':
        model = model_dict['bl9']
    elif m == '2':
        model = model_dict['dl10']
    elif m == '1' or m == '':
        model = model_dict['bl10']
    else:
        print('exit')

    name = input('Server Name >>>> ')
    SN = input('Serial Number >>>> ')
    site = input('Site (default Shanghai) >>>> ')    
    if site == '': site = 'Shanghai'
    # loc = input('eg:IDX2/A24/GSHFRB024A2/Bay7 >>> ')
    DC = input('DataCenter (default IDX2) >>>>')
    if DC == '': DC = 'IDX2'
    Rack = input('Rack required >>>>')
    if Rack == '' : raise 'Rack info missing'
    loc = f'{DC}/{Rack}'

    if m == '1' or m == '3':
        Enclosure = input('Enclosure >>>>')
        Bay = input('Bay number >>>>')
        if Enclosure != '' and Bay != '':
            loc = f'{loc}/{Enclosure}/Bay {Bay}'
        else:
            print('Invalid input')
    srv_dict = {
        'name':name,
        'SN':SN,
        'site':site,
        'loc':loc,
        'model':model
    }
    return srv_dict 


def batch():
    '''create SIR in batch'''
    xls = os.path.join(wp,'sir.xlsx')
    wb = openpyxl.load_workbook(xls)
    sheet = wb['hw']
    for x in range(2,sheet.max_row+1):
        if sheet.cell(row=x,column=11).value != 'Y':  # SIR mark Y
            site = sheet.cell(row=x,column=1).value
            DC = sheet.cell(row=x,column=2).value
            Rack = sheet.cell(row=x,column=3).value
            Enclosure = sheet.cell(row=x,column=4).value
            bay = sheet.cell(row=x,column=5).value
            m = sheet.cell(row=x,column=6).value  # model
            name = sheet.cell(row=x,column=7).value
            print(f'SIR {name}')
            SN = sheet.cell(row=x,column=8).value
            if m == "ProLiant DX360 Gen10":
                model = model_dict['nutanix']
                loc = f'{DC}/{Rack}'
            elif m == "HPE Synergy 480 Gen10":
                model = model_dict['s480g10']
                loc = f'{DC}/{Rack}/{Enclosure}/Bay{bay}'
            elif m == 'ProLiant BL460 Gen10':
                model = model_dict['460g10']
                loc = f'{DC}/{Rack}/{Enclosure}/Bay{bay}'
            elif m == 'ProLiant DL380 Gen10':
                model = model_dict['380g10']
                loc = f'{DC}/{Rack}'
            elif m == 'ProLiant BL460 Gen9':
                model = model_dict['460g9']
                loc = f'{DC}/{Rack}/{Enclosure}/Bay{bay}'
            elif m == 'ProLiant DL580 Gen10':
                model = model_dict['580g10']
                loc = f'{DC}/{Rack}' 
            else:
                print('invalid input')
                sys.exit()
            srv_dict = {
                'name':name,
                'SN':SN,
                'site':site,
                'loc':loc,
                'model':model
            }
            sir(srv_dict)
            sheet.cell(row=x,column=11).value = 'Y'
            try:
                wb.save(xls)
            except PermissionError as e:
                # print(e)
                print('Is file being opened?')


def main():
    parser = argparse.ArgumentParser(description = 'SIR tool')
    group = parser.add_mutually_exclusive_group()
    group.add_argument('-b','--batch',help='SIR batch mode',action='store_true')
    group.add_argument('-o','--one',help='SIR one mode',action='store_true')
    args = parser.parse_args()
    
    if args.batch:
        batch()
    elif args.one:
        srv_dict = select()
        sir(srv_dict)
    else:
        batch()

if __name__=='__main__':
    try:
        main()
    except KeyboardInterrupt:
        print('ctrl + c')

